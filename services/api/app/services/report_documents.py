from __future__ import annotations

from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

REPORT_ACCENT = "0F766E"
REPORT_DARK = "16343A"
REPORT_MUTED = "62757B"
REPORT_HEADER_FILL = "EAF6F3"
REPORT_LIGHT_FILL = "F5F8F8"
REPORT_BORDER = "C7D5D8"
PDF_FONT_REGULAR = "IELTSReportDeng"
PDF_FONT_BOLD = "IELTSReportDengBold"


def _text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    return str(value)


def _title(report: dict[str, Any], fallback: str) -> str:
    summary = report.get("summary") or {}
    assignment = report.get("assignment") or {}
    return _text(
        summary.get("title")
        or assignment.get("title")
        or fallback
        or "IELTS G类阅读教师报告"
    )


def _format_seconds(value: Any) -> str:
    seconds = max(0, int(value or 0))
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f"{hours}小时{minutes}分{seconds}秒"
    if minutes:
        return f"{minutes}分{seconds}秒"
    return f"{seconds}秒"


def _format_date(value: Any) -> str:
    text = _text(value, "暂无")
    return text[:10] if len(text) >= 10 else text


def _section_title(report: dict[str, Any]) -> str:
    return "作业模块与完成情况" if report.get("modules") else "练习记录与完成情况"


def _summary_metrics(report: dict[str, Any]) -> list[tuple[str, str]]:
    summary = report.get("summary") or {}
    return [
        ("首次练习", _text(summary.get("first_attempt_count"), "0")),
        ("累计正确率", f"{summary.get('accuracy', 0)}%"),
        ("总题数", _text(summary.get("total_questions"), "0")),
        ("可信总用时", _format_seconds(summary.get("total_elapsed_seconds"))),
        ("练习记录", _text(summary.get("session_count"), "0")),
        ("重做记录", _text(summary.get("retry_count"), "0")),
        ("答对题数", _text(summary.get("correct"), "0")),
        ("未作答", _text(summary.get("unanswered"), "0")),
    ]


def _module_rows(report: dict[str, Any]) -> list[list[str]]:
    modules = report.get("modules") or []
    if modules:
        return [
            [
                _text(item.get("title"), "未命名模块"),
                _text((item.get("progress") or {}).get("progress_text"), "暂无进度"),
                f"{item.get('score', 0)}/{item.get('total', 0)}",
                f"{item.get('accuracy', 0)}%",
                _format_seconds(item.get("trusted_seconds")),
                _text(item.get("unanswered"), "0"),
            ]
            for item in modules
        ]
    return [
        [
            _format_date(item.get("created_at")),
            _text(item.get("skill_label") or item.get("test_title"), "未命名练习"),
            "首次" if item.get("attempt_kind") == "first" else f"重做 {max(1, int(item.get('attempt_number') or 1) - 1)}",
            f"{item.get('score', 0)}/{item.get('total', 0)}",
            f"{item.get('accuracy', 0)}%",
            _format_seconds(item.get("elapsed_seconds")),
        ]
        for item in report.get("trend") or []
    ]


def _part_rows(report: dict[str, Any]) -> list[list[str]]:
    rows = report.get("part_results") or []
    return [
        [
            _text(item.get("title"), "未分类 Part"),
            _text(item.get("correct"), "0"),
            _text(item.get("total"), "0"),
            f"{item.get('accuracy', 0)}%",
            _text(item.get("status_label"), "数据不足"),
        ]
        for item in rows
    ]


def _type_rows(report: dict[str, Any]) -> list[list[str]]:
    cause_by_type: dict[str, str] = {}
    for question in report.get("representative_questions") or []:
        question_type = _text(question.get("question_type"), "未分类")
        if question_type not in cause_by_type and question.get("cause_label"):
            cause_by_type[question_type] = _text(question.get("cause_label"))
    return [
        [
            _text(item.get("question_type") or item.get("question_subtype"), "未分类"),
            f"{item.get('correct', 0)}/{item.get('total', 0)}",
            f"{item.get('accuracy', 0)}%",
            _text(item.get("status_label"), "数据不足"),
            cause_by_type.get(
                _text(item.get("question_type"), "未分类"),
                "需结合代表错题确认",
            ),
        ]
        for item in report.get("question_type_matrix") or []
    ]


def _cause_rows(report: dict[str, Any]) -> list[list[str]]:
    return [
        [
            _text(item.get("label"), "未分类"),
            _text(item.get("count"), "0"),
            _text(item.get("session_count"), "0"),
            "；".join(_text(value) for value in item.get("examples") or []) or "暂无",
        ]
        for item in report.get("error_cause_distribution") or []
    ]


def _set_docx_font(
    run: Any,
    *,
    size: float = 10.5,
    bold: bool = False,
    color: str = REPORT_DARK,
    italic: bool = False,
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_cell_fill(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, *, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _configure_docx_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for name, size, before, after, color in (
        ("Heading 1", 15, 12, 6, REPORT_ACCENT),
        ("Heading 2", 12, 9, 4, "1F4D78"),
        ("Heading 3", 11, 7, 3, REPORT_DARK),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _docx_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    color: str = REPORT_DARK,
    size: float = 10.5,
    italic: bool = False,
    after: float = 6,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    _set_docx_font(
        paragraph.add_run(text),
        size=size,
        bold=bold,
        color=color,
        italic=italic,
    )
    return paragraph


def _docx_bullets(document: Document, values: Iterable[str]) -> None:
    for value in values:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        _set_docx_font(paragraph.add_run(_text(value)))


def _docx_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_mm: list[float],
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Mm(widths_mm[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_fill(cell, REPORT_HEADER_FILL)
        _set_cell_margins(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_docx_font(cell.paragraphs[0].add_run(header), size=9.5, bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Mm(widths_mm[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cells[index])
            if row_index % 2:
                _set_cell_fill(cells[index], REPORT_LIGHT_FILL)
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            )
            _set_docx_font(paragraph.add_run(_text(value)), size=9.2)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_docx_font(run, size=8.5, color=REPORT_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    suffix = paragraph.add_run(" 页")
    _set_docx_font(suffix, size=8.5, color=REPORT_MUTED)


def _add_docx_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_docx_font(
        header.add_run("IELTS G类阅读 AI 教练 · 真人老师教学诊断报告"),
        size=8.5,
        color=REPORT_MUTED,
    )
    footer = section.footer.paragraphs[0]
    footer.text = ""
    _add_page_field(footer)


def _add_docx_cover(document: Document, report: dict[str, Any], title: str) -> None:
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(6)
    _set_docx_font(
        kicker.add_run("IELTS GENERAL TRAINING READING"),
        size=9,
        bold=True,
        color=REPORT_ACCENT,
    )
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(5)
    _set_docx_font(
        heading.add_run(_title(report, title)),
        size=20,
        bold=True,
        color=REPORT_DARK,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    _set_docx_font(
        subtitle.add_run(
            f"{report.get('layout_label') or '确定性学习报告'} · "
            f"报告引擎 {report.get('engine_version') or '—'}"
        ),
        size=9,
        color=REPORT_MUTED,
    )
    summary = report.get("summary") or {}
    metadata = (
        f"统计周期：{_format_date(summary.get('date_from'))} - {_format_date(summary.get('date_to'))}"
        f"    报告来源：{report.get('generated_from') or 'persisted_sessions'}"
    )
    meta = _docx_paragraph(document, metadata, size=9, color=REPORT_MUTED, after=10)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    assignment = report.get("assignment") or {}
    if assignment.get("description"):
        _docx_paragraph(
            document,
            f"作业说明：{assignment['description']}",
            color=REPORT_DARK,
            after=8,
        )


def _add_docx_summary(document: Document, report: dict[str, Any]) -> None:
    metrics = _summary_metrics(report)
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [28.0, 64.0, 28.0, 64.0]
    for row_index in range(4):
        left_label, left_value = metrics[row_index * 2]
        right_label, right_value = metrics[row_index * 2 + 1]
        values = (left_label, left_value, right_label, right_value)
        for column, value in enumerate(values):
            cell = table.rows[row_index].cells[column]
            cell.width = Mm(widths[column])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=100, bottom=100)
            if column in (0, 2):
                _set_cell_fill(cell, REPORT_HEADER_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_docx_font(
                paragraph.add_run(value),
                size=9.5 if column in (0, 2) else 11,
                bold=column not in (0, 2),
                color=REPORT_MUTED if column in (0, 2) else REPORT_DARK,
            )
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_docx_questions(document: Document, report: dict[str, Any]) -> None:
    questions = report.get("representative_questions") or []
    if not questions:
        _docx_paragraph(document, "本范围没有可展示的错题。")
        return
    for item in questions:
        source = _text(
            item.get("source"),
            f"{item.get('test_title') or '练习'} / Q{item.get('question_number')}",
        )
        _docx_paragraph(document, source, bold=True, color=REPORT_ACCENT, after=2)
        _docx_paragraph(
            document,
            f"题型：{item.get('question_type') or '未分类'}    "
            f"错因：{item.get('cause_label') or '待确认'}",
            size=9.5,
            color=REPORT_MUTED,
            after=3,
        )
        _docx_paragraph(document, f"题干：{item.get('prompt') or '—'}", bold=True, after=3)
        _docx_paragraph(
            document,
            f"学生答案：{item.get('user_answer') or '未作答'}    "
            f"正确答案：{item.get('correct_answer') or '—'}",
            after=3,
        )
        evidence = "；".join(_text(value) for value in item.get("evidence") or [])
        _docx_paragraph(
            document,
            f"原文定位：{evidence or '当前题库未提供可核验定位句'}",
            after=3,
        )
        if item.get("analysis"):
            _docx_paragraph(document, f"答案解析：{item['analysis']}", after=3)
        _docx_paragraph(
            document,
            f"学生确认：{item.get('student_confirmation_label') or '未记录'}",
            after=3,
        )
        _docx_paragraph(
            document,
            f"请老师重点观察：{item.get('teacher_observation') or '复核学生的定位和判断过程。'}",
            italic=True,
            color="1F4D78",
            after=8,
        )


def build_teacher_docx(report: dict[str, Any], *, title: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    _configure_docx_styles(document)
    _add_docx_header_footer(document)
    _add_docx_cover(document, report, title)
    _add_docx_summary(document, report)

    document.add_heading("1. 给老师的核心摘要", level=1)
    _docx_bullets(
        document,
        report.get("deterministic_interpretation")
        or ["当前样本不足，暂不作能力定性。"],
    )

    document.add_heading(f"2. {_section_title(report)}", level=1)
    _docx_table(
        document,
        (
            ["模块", "完成情况", "成绩", "正确率", "可信用时", "未作答"]
            if report.get("modules")
            else ["日期", "练习", "性质", "成绩", "正确率", "用时"]
        ),
        _module_rows(report),
        [40, 48, 24, 24, 28, 20] if report.get("modules") else [24, 58, 22, 24, 24, 32],
    )

    document.add_heading("3. Part 与练习表现", level=1)
    if _part_rows(report):
        _docx_table(
            document,
            ["Part", "正确", "总数", "正确率", "判断"],
            _part_rows(report),
            [55, 24, 24, 32, 49],
        )
    else:
        _docx_paragraph(document, "当前记录没有可单独汇总的 Part 数据。")

    document.add_heading("4. 总体题型能力矩阵", level=1)
    _docx_table(
        document,
        ["题型", "正确", "正确率", "状态", "主要错因"],
        _type_rows(report),
        [58, 24, 28, 30, 44],
    )

    document.add_heading("5. 总体错误原因分布", level=1)
    if _cause_rows(report):
        _docx_table(
            document,
            ["错误原因", "题数", "涉及练习", "示例"],
            _cause_rows(report),
            [58, 22, 30, 74],
        )
    else:
        _docx_paragraph(document, "当前范围没有可统计的错误原因。")

    document.add_heading("6. 代表性错题", level=1)
    _add_docx_questions(document, report)

    document.add_heading("7. 给老师的教学参考", level=1)
    _docx_bullets(
        document,
        report.get("teacher_observation_points")
        or ["结合学生实际作答过程复核定位、排除和最终确认步骤。"],
    )

    document.add_heading("8. 数据口径", level=1)
    _docx_bullets(document, report.get("data_notes") or [])
    _docx_paragraph(
        document,
        "本报告不调用 AI；标准答案、判分和 Band 规则没有改变。最终教学判断和课堂安排由真人老师作出。",
        bold=True,
        color=REPORT_ACCENT,
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_font_names() -> tuple[str, str]:
    registered = set(pdfmetrics.getRegisteredFontNames())
    if PDF_FONT_REGULAR in registered and PDF_FONT_BOLD in registered:
        return PDF_FONT_REGULAR, PDF_FONT_BOLD
    regular_path = Path("C:/Windows/Fonts/Deng.ttf")
    bold_path = Path("C:/Windows/Fonts/Dengb.ttf")
    if regular_path.exists() and bold_path.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT_REGULAR, str(regular_path)))
        pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(bold_path)))
        return PDF_FONT_REGULAR, PDF_FONT_BOLD
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except KeyError:
        pass
    return "STSong-Light", "STSong-Light"


def _pdf_styles() -> dict[str, ParagraphStyle]:
    regular_font, bold_font = _pdf_font_names()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName=bold_font,
            fontSize=20,
            leading=27,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{REPORT_DARK}"),
            spaceAfter=5,
        ),
        "subtitle": ParagraphStyle(
            "ReportSubtitle",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=8.5,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{REPORT_MUTED}"),
            spaceAfter=12,
        ),
        "heading": ParagraphStyle(
            "ReportHeading",
            parent=base["Heading2"],
            fontName=bold_font,
            fontSize=13,
            leading=18,
            textColor=colors.HexColor(f"#{REPORT_ACCENT}"),
            spaceBefore=10,
            spaceAfter=7,
        ),
        "subheading": ParagraphStyle(
            "ReportSubheading",
            parent=base["Heading3"],
            fontName=bold_font,
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=5,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=base["BodyText"],
            fontName=regular_font,
            fontSize=9.2,
            leading=14,
            textColor=colors.HexColor(f"#{REPORT_DARK}"),
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "ReportSmall",
            parent=base["BodyText"],
            fontName=regular_font,
            fontSize=7.6,
            leading=11,
            textColor=colors.HexColor(f"#{REPORT_DARK}"),
        ),
        "muted": ParagraphStyle(
            "ReportMuted",
            parent=base["BodyText"],
            fontName=regular_font,
            fontSize=7.5,
            leading=11,
            textColor=colors.HexColor(f"#{REPORT_MUTED}"),
            spaceAfter=3,
        ),
    }


def _p(text: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(_text(text)).replace("\n", "<br/>"), style)


def _pdf_table(
    headers: list[str],
    rows: list[list[str]],
    styles: dict[str, ParagraphStyle],
    widths_mm: list[float],
) -> Table:
    regular_font, bold_font = _pdf_font_names()
    data = [[_p(value, styles["small"]) for value in headers]]
    data.extend([[_p(value, styles["small"]) for value in row] for row in rows])
    table = Table(
        data,
        colWidths=[value * mm for value in widths_mm],
        repeatRows=1,
        hAlign="CENTER",
    )
    commands: list[tuple[Any, ...]] = [
        ("FONTNAME", (0, 0), (-1, -1), regular_font),
        ("FONTNAME", (0, 0), (-1, 0), bold_font),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{REPORT_HEADER_FILL}")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{REPORT_BORDER}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (-2, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if len(data) > 2:
        commands.append(
            ("BACKGROUND", (0, 2), (-1, -1), colors.HexColor(f"#{REPORT_LIGHT_FILL}"))
        )
    table.setStyle(
        TableStyle(commands)
    )
    return table


def _pdf_header_footer(canvas: Any, document: Any) -> None:
    regular_font, _ = _pdf_font_names()
    canvas.saveState()
    canvas.setFont(regular_font, 7.5)
    canvas.setFillColor(colors.HexColor(f"#{REPORT_MUTED}"))
    canvas.drawString(13 * mm, 287 * mm, "IELTS G类阅读 AI 教练 · 真人老师教学诊断报告")
    canvas.drawRightString(197 * mm, 8 * mm, f"第 {document.page} 页")
    canvas.restoreState()


def _pdf_summary_table(report: dict[str, Any], styles: dict[str, ParagraphStyle]) -> Table:
    regular_font, bold_font = _pdf_font_names()
    metrics = _summary_metrics(report)
    data: list[list[Paragraph]] = []
    for index in range(0, len(metrics), 2):
        left, right = metrics[index], metrics[index + 1]
        data.append(
            [
                _p(left[0], styles["muted"]),
                _p(left[1], styles["body"]),
                _p(right[0], styles["muted"]),
                _p(right[1], styles["body"]),
            ]
        )
    table = Table(data, colWidths=[28 * mm, 64 * mm, 28 * mm, 64 * mm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), regular_font),
                ("FONTNAME", (0, 0), (0, -1), bold_font),
                ("FONTNAME", (2, 0), (2, -1), bold_font),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{REPORT_HEADER_FILL}")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor(f"#{REPORT_HEADER_FILL}")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{REPORT_BORDER}")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _pdf_questions(report: dict[str, Any], styles: dict[str, ParagraphStyle]) -> list[Any]:
    story: list[Any] = []
    questions = report.get("representative_questions") or []
    if not questions:
        return [_p("本范围没有可展示的错题。", styles["body"])]
    for item in questions:
        source = _text(
            item.get("source"),
            f"{item.get('test_title') or '练习'} / Q{item.get('question_number')}",
        )
        evidence = "；".join(_text(value) for value in item.get("evidence") or [])
        elements = [
            _p(source, styles["subheading"]),
            _p(
                f"题型：{item.get('question_type') or '未分类'}　"
                f"错因：{item.get('cause_label') or '待确认'}",
                styles["muted"],
            ),
            _p(f"题干：{item.get('prompt') or '—'}", styles["body"]),
            _p(
                f"学生答案：{item.get('user_answer') or '未作答'}　"
                f"正确答案：{item.get('correct_answer') or '—'}",
                styles["body"],
            ),
            _p(
                f"原文定位：{evidence or '当前题库未提供可核验定位句'}",
                styles["body"],
            ),
        ]
        if item.get("analysis"):
            elements.append(_p(f"答案解析：{item['analysis']}", styles["body"]))
        elements.extend(
            [
                _p(
                    f"学生确认：{item.get('student_confirmation_label') or '未记录'}",
                    styles["muted"],
                ),
                _p(
                    f"请老师重点观察：{item.get('teacher_observation') or '复核学生的定位和判断过程。'}",
                    styles["body"],
                ),
                Spacer(1, 4),
            ]
        )
        story.append(KeepTogether(elements))
    return story


def build_teacher_pdf(report: dict[str, Any], *, title: str) -> bytes:
    buffer = BytesIO()
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=13 * mm,
        leftMargin=13 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title=_title(report, title),
        author="IELTS G类阅读 AI 教练",
        subject="真人老师教学诊断报告",
    )
    summary = report.get("summary") or {}
    assignment = report.get("assignment") or {}
    story: list[Any] = [
        _p("IELTS GENERAL TRAINING READING", styles["subtitle"]),
        _p(_title(report, title), styles["title"]),
        _p(
            f"{report.get('layout_label') or '确定性学习报告'} · "
            f"统计周期：{_format_date(summary.get('date_from'))} - {_format_date(summary.get('date_to'))} · "
            f"报告引擎 {report.get('engine_version') or '—'}",
            styles["subtitle"],
        ),
    ]
    if assignment.get("description"):
        story.append(_p(f"作业说明：{assignment['description']}", styles["body"]))
    story.extend([_pdf_summary_table(report, styles), Spacer(1, 6)])

    story.append(_p("1. 给老师的核心摘要", styles["heading"]))
    for value in report.get("deterministic_interpretation") or ["当前样本不足，暂不作能力定性。"]:
        story.append(_p(f"- {value}", styles["body"]))

    story.append(_p(f"2. {_section_title(report)}", styles["heading"]))
    story.append(
        _pdf_table(
            (
                ["模块", "完成情况", "成绩", "正确率", "可信用时", "未作答"]
                if report.get("modules")
                else ["日期", "练习", "性质", "成绩", "正确率", "用时"]
            ),
            _module_rows(report),
            styles,
            [40, 48, 24, 24, 28, 20] if report.get("modules") else [24, 58, 22, 24, 24, 32],
        )
    )

    story.append(_p("3. Part 与练习表现", styles["heading"]))
    if _part_rows(report):
        story.append(
            _pdf_table(
                ["Part", "正确", "总数", "正确率", "判断"],
                _part_rows(report),
                styles,
                [55, 24, 24, 32, 49],
            )
        )
    else:
        story.append(_p("当前记录没有可单独汇总的 Part 数据。", styles["body"]))

    story.extend(
        [
            _p("4. 总体题型能力矩阵", styles["heading"]),
            _pdf_table(
                ["题型", "正确", "正确率", "状态", "主要错因"],
                _type_rows(report),
                styles,
                [58, 24, 28, 30, 44],
            ),
            _p("5. 总体错误原因分布", styles["heading"]),
        ]
    )
    if _cause_rows(report):
        story.append(
            _pdf_table(
                ["错误原因", "题数", "涉及练习", "示例"],
                _cause_rows(report),
                styles,
                [58, 22, 30, 74],
            )
        )
    else:
        story.append(_p("当前范围没有可统计的错误原因。", styles["body"]))

    story.extend([PageBreak(), _p("6. 代表性错题", styles["heading"])])
    story.extend(_pdf_questions(report, styles))

    story.append(_p("7. 给老师的教学参考", styles["heading"]))
    for value in report.get("teacher_observation_points") or [
        "结合学生实际作答过程复核定位、排除和最终确认步骤。"
    ]:
        story.append(_p(f"- {value}", styles["body"]))

    story.append(_p("8. 数据口径", styles["heading"]))
    for value in report.get("data_notes") or []:
        story.append(_p(f"- {value}", styles["small"]))
    story.append(
        _p(
            "本报告不调用 AI；标准答案、判分和 Band 规则没有改变。最终教学判断和课堂安排由真人老师作出。",
            styles["body"],
        )
    )
    document.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    return buffer.getvalue()
